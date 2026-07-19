import streamlit as st
import os
import json
import hashlib
import time  # 引入時間模組，用來控制發送頻率
from google import genai
from google.genai import types
from pydantic import BaseModel
import fitz  # PyMuPDF
from PIL import Image
import io
from docx import Document
from docx.shared import Inches, Pt

# --- 1. 定義嚴格的 JSON Schema ---
class VocabularyItem(BaseModel):
    word: str
    reading: str
    meaning: str

class PageAnalysis(BaseModel):
    summary: str
    vocabulary: list[VocabularyItem]

# --- 2. 設定網頁外觀 ---
st.set_page_config(page_title="小學日文課本翻譯機 (快取排版版)", layout="wide", page_icon="📚")
st.title("📚 小學日文課本翻譯機 (快取與左右排版版)")
st.write("已啟用「本地快取」與「防 429 斷線重試機制」，安心處理多頁 PDF！")

# --- 3. 檢查 API Key ---
api_key = os.environ.get("GEMINI_API_KEY")
if not api_key:
    st.error("⚠️ 找不到 GEMINI_API_KEY 環境變數。請確認設定。")
    st.stop() 

# --- 4. 檔案上傳區塊 ---
uploaded_file = st.file_uploader("請上傳一份 PDF 格式的課本檔案", type=["pdf"])

if uploaded_file is not None:
    if st.button("開始解析 (支援快取) ✨"):
        try:
            client = genai.Client(api_key=api_key)
            pdf_bytes = uploaded_file.read()
            
            # --- 利用 PDF 內容產生 MD5 Hash，建立快取機制 ---
            pdf_hash = hashlib.md5(pdf_bytes).hexdigest()
            cache_dir = "cache_data"
            os.makedirs(cache_dir, exist_ok=True)
            
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            total_pages = len(pdf_document)
            
            st.success(f"成功讀取，共 {total_pages} 頁。檔案 Hash: {pdf_hash[:8]}...")

            # 初始化 Word 文件
            doc = Document()
            doc.add_heading('📚 小學日文課本解析筆記', 0)

            prompt = '''
            你是一位專業的日本小學老師。請閱讀這頁課本圖片。
            1. 寫出繁體中文大意總結 (適合小學生閱讀)。
            2. 挑選 5 到 10 個最重要的單字。
            單字的 reading (念法) 請嚴格使用『片假名』標示。
            '''

            for page_num in range(total_pages):
                st.subheader(f"📖 第 {page_num + 1} 頁")
                
                # 準備快取檔案路徑
                cache_file = os.path.join(cache_dir, f"{pdf_hash}_page_{page_num}.json")
                
                # 處理圖片
                page = pdf_document.load_page(page_num)
                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))
                
                # Streamlit 左右排版顯示
                col1, col2 = st.columns([1, 1])
                with col1:
                    st.image(image, caption=f"第 {page_num + 1} 頁", use_column_width=True)

                data = None
                
                # --- 檢查是否有快取 ---
                if os.path.exists(cache_file):
                    with open(cache_file, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    st.info("⚡ 命中快取 (Cache Hit)！略過 API 呼叫。")
                else:
                    # --- 加入重試機制 (Retry Logic) ---
                    success = False
                    while not success:
                        try:
                            with st.spinner(f'呼叫 API 處理第 {page_num + 1} 頁中...'):
                                response = client.models.generate_content(
                                    model='gemini-3.5-flash',
                                    contents=[image, prompt],
                                    config=types.GenerateContentConfig(
                                        response_mime_type="application/json",
                                        response_schema=PageAnalysis,
                                        temperature=0.1 
                                    )
                                )
                            data = json.loads(response.text)
                            
                            # 儲存到快取
                            with open(cache_file, "w", encoding="utf-8") as f:
                                json.dump(data, f, ensure_ascii=False, indent=2)
                            
                            success = True # 標記成功，跳出 while 迴圈
                            
                            # 為了避免連續呼叫撞到 429 限制，成功後強制暫停 5 秒
                            if page_num < total_pages - 1:
                                time.sleep(5)
                                
                        except Exception as e:
                            error_msg = str(e)
                            # 如果錯誤訊息包含 429 或 RESOURCE_EXHAUSTED，就進入等待
                            if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                                st.warning("⚠️ 速度太快觸發免費額度限制，程式將自動暫停 60 秒後接續重試，請稍候...")
                                time.sleep(60)
                            else:
                                raise e # 如果是其他未知的錯誤，直接報錯停止

                # 網頁上的資料顯示
                with col2:
                    st.markdown(f"**📝 大意：** {data.get('summary', '無')}")
                    vocab_md = "| 日文 | 念法 | 中文 |\n|---|---|---|\n"
                    for v in data.get('vocabulary', []):
                        vocab_md += f"| {v.get('word','')} | {v.get('reading','')} | {v.get('meaning','')} |\n"
                    st.markdown(vocab_md)
                
                st.markdown("---")

                # --- 寫入 Word 左右排版 ---
                doc.add_heading(f'第 {page_num + 1} 頁', level=1)
                layout_table = doc.add_table(rows=1, cols=2)
                layout_table.autofit = False
                layout_table.columns[0].width = Inches(3.5)
                layout_table.columns[1].width = Inches(3.5)
                
                left_cell = layout_table.rows[0].cells[0]
                right_cell = layout_table.rows[0].cells[1]
                
                # 左側圖片
                left_p = left_cell.paragraphs[0]
                left_run = left_p.add_run()
                img_stream = io.BytesIO(img_data)
                left_run.add_picture(img_stream, width=Inches(3.2)) 
                
                # 右側文字
                right_p = right_cell.paragraphs[0]
                right_p.add_run("📝 大意：\n").bold = True
                right_p.add_run(data.get('summary', '') + "\n\n")
                right_p.add_run("💡 重點單字：").bold = True
                
                vocab_list = data.get('vocabulary', [])
                if vocab_list:
                    vocab_table = right_cell.add_table(rows=1, cols=3)
                    vocab_table.style = 'Table Grid'
                    
                    hdr_cells = vocab_table.rows[0].cells
                    hdr_cells[0].text = '日文'
                    hdr_cells[1].text = '念法'
                    hdr_cells[2].text = '中文意思'
                    
                    for v in vocab_list:
                        row_cells = vocab_table.add_row().cells
                        row_cells[0].text = v.get('word', '')
                        row_cells[1].text = v.get('reading', '')
                        row_cells[2].text = v.get('meaning', '')
                        
                if page_num < total_pages - 1:
                    doc.add_page_break() 

            st.balloons()
            st.success("🎉 整份 PDF 處理完成！請點擊下方按鈕下載左右排版筆記。")
            
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            
            st.download_button(
                label="📥 下載 Word 左右排版筆記 (可上傳至 Google Docs)",
                data=doc_stream,
                file_name="日文課本_排版筆記.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            pdf_document.close()

        except Exception as e:
            st.error(f"過程中發生錯誤：{e}")
